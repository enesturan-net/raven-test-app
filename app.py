import streamlit as st
from datetime import date, datetime
from docx import Document
from docx.shared import Pt
import io
import base64
import random
import os

# --- SAYFA AYARLARI ---
st.set_page_config(page_title="Çocuk Raven Test Analizi", layout="centered", page_icon="🧠")

# --- HAFIZA AYARLARI ---
if 'analiz_yapildi' not in st.session_state: st.session_state['analiz_yapildi'] = False
if 'sonuclar' not in st.session_state: st.session_state['sonuclar'] = []
if 'popup_ac' not in st.session_state: st.session_state['popup_ac'] = False
if 'kisi_bilgi' not in st.session_state: st.session_state['kisi_bilgi'] = {}

def popup_tetikle(): st.session_state['popup_ac'] = True

# --- PUAN DÖNÜŞTÜRME (KONTROL NOKTASI 1) ---
# Tablodaki aralıkların (örn: 59-60) ÜST değerlerini aldım.
# Hata buradaysa bu sayıları değiştirebiliriz.
def puani_donustur(p):
    mapping = {
        28: 60, 27: 58, 26: 57, 25: 56, 24: 55, 23: 54, 22: 53,
        21: 52, 20: 51, 19: 50, 18: 48, 17: 47, 16: 46, 15: 45,
        14: 44, 13: 42, 12: 41, 11: 40, 10: 39, 9: 37, 8: 35,
        7: 34, 6: 32, 5: 30, 4: 27, 3: 24, 2: 20, 1: 16, 0: 0
    }
    return mapping.get(p, 0)

ulke_isimleri = {
    "UK": "İngiltere", "US": "ABD", "CN": "Çin", "NZ": "Yeni Zelanda", 
    "AU": "Avustralya", "PL": "Polonya", "SI": "Slovenya", "AR": "Arjantin", 
    "QA": "Katar", "NL": "Hollanda", "FR": "Fransa", "TW": "Tayvan", 
    "SK": "Slovakya", "CH": "İsviçre", "RU": "Rusya"
}

# --- VERİ TABANI (KONTROL NOKTASI 2) ---
# Buradaki değerler PDF'ten alındı. Yanlışlık varsa burayı düzelteceğiz.
veritabani = {
    "UK": {
        "75-80": {95:33, 90:30, 75:22, 50:16, 25:13}, "81-86": {95:34, 90:32, 75:26, 50:19, 25:14},
        "87-92": {95:37, 90:35, 75:30, 50:22, 25:15}, "93-98": {95:40, 90:38, 75:33, 50:25, 25:17},
        "99-104":{95:42, 90:40, 75:36, 50:31, 25:22}, "105-110":{95:44, 90:42, 75:38, 50:33, 25:25},
        "111-116":{95:46, 90:44, 75:41, 50:36, 25:28}, "117-122":{95:48, 90:46, 75:42, 50:38, 25:32},
        "123-128":{95:49, 90:47, 75:43, 50:39, 25:33}, "129-134":{95:50, 90:48, 75:44, 50:40, 25:34},
        "135-140":{95:51, 90:49, 75:45, 50:41, 25:36}, "141-146":{95:52, 90:50, 75:46, 50:41, 25:37},
        "147-152":{95:53, 90:51, 75:47, 50:42, 25:38}, "153-158":{95:54, 90:52, 75:49, 50:43, 25:39},
        "159-164":{95:54, 90:53, 75:49, 50:44, 25:41}, "165-186":{95:55, 90:54, 75:50, 50:45, 25:42},
        "204-275": {95:59, 90:58, 75:57, 50:54, 25:49}, "276-335": {95:59, 90:58, 75:57, 50:54, 25:49},
        "336-395": {95:59, 90:58, 75:57, 50:54, 25:49}, "396-455": {95:59, 90:58, 75:56, 50:54, 25:49},
        "456-515": {95:59, 90:58, 75:56, 50:53, 25:48},
    },
    "US": {
        "78-83": {95:30, 90:27, 75:21, 50:14, 25:12}, "84-89": {95:33, 90:30, 75:25, 50:17, 25:13},
        "90-95": {95:36, 90:33, 75:28, 50:20, 25:14}, "96-101":{95:38, 90:36, 75:31, 50:23, 25:16},
        "102-107":{95:40, 90:38, 75:34, 50:26, 25:18}, "108-113":{95:42, 90:40, 75:36, 50:29, 25:21},
        "114-119":{95:44, 90:42, 75:38, 50:32, 25:24}, "120-125":{95:46, 90:44, 75:41, 50:34, 25:26},
        "126-131":{95:47, 90:45, 75:41, 50:36, 25:28}, "132-137":{95:48, 90:46, 75:43, 50:37, 25:30},
        "138-143":{95:49, 90:47, 75:44, 50:38, 25:32}, "144-149":{95:50, 90:48, 75:45, 50:39, 25:33},
        "150-198":{95:51, 90:49, 75:46, 50:40, 25:34},
        "204-275": {95:59, 90:58, 75:56, 50:52, 25:47}, "276-335": {95:59, 90:58, 75:56, 50:52, 25:47},
        "336-395": {95:59, 90:58, 75:56, 50:52, 25:47}, "396-455": {95:59, 90:58, 75:56, 50:52, 25:47},
    },
    "PL": {
        "72-77": {95:26, 90:23, 75:19, 50:14, 25:13}, "78-83": {95:33, 90:30, 75:22, 50:16, 25:13},
        "84-89": {95:29, 90:25, 75:20, 50:16, 25:14}, "90-95": {95:34, 90:32, 75:26, 50:19, 25:14},
        "96-101": {95:31, 90:27, 75:21, 50:17, 25:14}, "102-107":{95:37, 90:35, 75:30, 50:22, 25:15},
        "108-113":{95:35, 90:29, 75:23, 50:18, 25:15}, "114-119":{95:40, 90:38, 75:33, 50:25, 25:17},
        "120-125":{95:42, 90:40, 75:36, 50:31, 25:22}, "132-143":{95:50, 90:47, 75:43, 50:39, 25:33},
        "144-155":{95:51, 90:50, 75:45, 50:41, 25:37}, "156-167":{95:54, 90:52, 75:48, 50:43, 25:40},
        "168-179":{95:54, 90:53, 75:49, 50:44, 25:41}, "180-203":{95:57, 90:55, 75:51, 50:47, 25:42},
        "204-215":{95:55, 90:53, 75:50, 50:47, 25:42}, "216-233":{95:57, 90:55, 75:52, 50:49, 25:44},
        "234-257":{95:59, 90:58, 75:57, 50:54, 25:49}, "258-293":{95:59, 90:58, 75:57, 50:54, 25:49},
        "294-329":{95:59, 90:58, 75:57, 50:54, 25:49}, "330-389":{95:59, 90:58, 75:56, 50:54, 25:48},
        "390-509":{95:56, 90:53, 75:48, 50:41, 25:36}, "510-629":{95:59, 90:58, 75:56, 50:52, 25:47},
    },
    # (Diğer ülkeler aynen kalsın...)
    "CN": {"63-74": {95:34, 90:29, 75:25, 50:16, 25:13}, "199-228":{95:58, 90:56, 75:53, 50:49, 25:45}, "441-560":{95:57, 90:56, 75:54, 50:48, 25:44}},
    "QA": {"69-80": {95:19, 90:18, 75:15, 50:14, 25:11}, "129-142":{95:44, 90:41, 75:34, 50:26, 25:19}}
}

# --- TASARIM ---
st.markdown("""
<style>
    .stApp { background-color: #ffffff; }
    h1, h2, h3, p, div, span, label { color: #262730 !important; }
    .stTextInput input, .stNumberInput input, .stDateInput input { background-color: #f0f2f6 !important; color: #000 !important; border: 1px solid #d6d6d6; }
    div.stButton > button { background-color: #FF4B4B; color: white !important; width: 100%; }
    .debug-box { background-color: #f8f9fa; padding: 15px; border-radius: 10px; border-left: 5px solid #ff4b4b; margin-top: 20px; font-size: 14px; }
</style>
""", unsafe_allow_html=True)

st.title("Raven Testi: Otomatik Eşleştirme")

col1, col2 = st.columns(2)
with col1: ad_soyad = st.text_input("Ad Soyad")
with col2: dob = st.date_input("Doğum Tarihi", min_value=date(1920, 1, 1))
dogru = st.number_input("Test Doğru Sayısı (0-28 Arası)", min_value=0, max_value=28, step=1)

# --- ANALİZ BUTONU ---
if st.button("Analiz Et", type="primary"):
    if not ad_soyad:
        st.error("Lütfen Ad Soyad giriniz.")
    else:
        bugun = date.today()
        yas_ay_toplam = (bugun.year - dob.year) * 12 + (bugun.month - dob.month)
        if bugun.day < dob.day: yas_ay_toplam -= 1
        
        yas_yil = yas_ay_toplam // 12
        yas_ay_artik = yas_ay_toplam % 12
        spm_puani = puani_donustur(dogru)

        st.session_state['analiz_yapildi'] = True
        st.session_state['kisi_bilgi'] = {
            "ad": ad_soyad, "dob": dob, "yas_yil": yas_yil, 
            "yas_ay": yas_ay_artik, "dogru": dogru, "spm": spm_puani, "toplam_ay": yas_ay_toplam
        }
        
        # Hesaplama ve Debug Verisi Toplama
        temp_sonuclar = []
        debug_log = []
        
        for ulke_kodu, veri in veritabani.items():
            ulke_adi = ulke_isimleri.get(ulke_kodu, ulke_kodu)
            bulunan_aralik = None
            for aralik_key in veri:
                min_ay, max_ay = map(int, aralik_key.split("-"))
                if min_ay <= yas_ay_toplam <= max_ay:
                    bulunan_aralik = veri[aralik_key]
                    debug_log.append(f"✅ {ulke_adi}: {min_ay}-{max_ay} ay aralığı bulundu. (Eşikler: {bulunan_aralik})")
                    break
            
            if bulunan_aralik:
                yuzdelik_sonuc = "5. Altı"
                dilimler = sorted(bulunan_aralik.keys(), reverse=True)
                for dilim in dilimler:
                    if spm_puani >= bulunan_aralik[dilim]:
                        yuzdelik_sonuc = f"%{dilim}"
                        break
                temp_sonuclar.append((ulke_adi, yuzdelik_sonuc))
            else:
                debug_log.append(f"❌ {ulke_adi}: {yas_ay_toplam} ay için uygun aralık YOK.")
        
        st.session_state['sonuclar'] = temp_sonuclar
        st.session_state['debug_log'] = debug_log

# --- SONUÇ EKRANI ---
if st.session_state['analiz_yapildi']:
    bilgi = st.session_state['kisi_bilgi']
    st.success(f"Hesaplama Tamamlandı")
    
    # --- DENETİM / DEBUG ALANI (HATAYI BULMAK İÇİN) ---
    with st.expander("🛠️ HESAPLAMA DETAYLARINI GÖSTER (Hata Kontrolü)"):
        st.markdown(f"""
        **1. Yaş Hesabı:**
        - Doğum Tarihi: {bilgi['dob']}
        - Bugün: {date.today()}
        - **Hesaplanan Toplam Ay:** {bilgi['toplam_ay']} Ay ({bilgi['yas_yil']} Yıl {bilgi['yas_ay']} Ay)
        
        **2. Puan Dönüşümü:**
        - Girilen Doğru (0-28): **{bilgi['dogru']}**
        - SPM Puanı (0-60): **{bilgi['spm']}** (Bu puan tabloda aranıyor)
        
        **3. Ülke Bazlı Kontrol Logları:**
        """)
        for log in st.session_state.get('debug_log', []):
            st.text(log)

    st.subheader("Sonuçlar")
    if not st.session_state['sonuclar']:
        st.warning("Veri bulunamadı.")
    else:
        for ulke, yuzdelik in st.session_state['sonuclar']:
            st.write(f"**{ulke}:** {yuzdelik}")
        
        # Word İndirme
        doc = Document()
        doc.add_heading('RAVEN TESTİ RAPORU', 0)
        p = doc.add_paragraph()
        p.add_run(f"Ad: {bilgi['ad']}\n").bold = True
        p.add_run(f"Yaş: {bilgi['yas_yil']} Yıl {bilgi['yas_ay']} Ay\n")
        p.add_run(f"Puan: {bilgi['spm']} (Ham: {bilgi['dogru']})")
        for ulke, yuzdelik in st.session_state['sonuclar']:
            doc.add_paragraph(f"{ulke}: {yuzdelik}", style='List Bullet')
        
        bio = io.BytesIO()
        doc.save(bio)
        st.download_button("📥 Raporu İndir", bio.getvalue(), f"Rapor_{bilgi['ad']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary", on_click=popup_tetikle)

@st.dialog("⚠️ UYARI")
def show_popup_modal():
    image_path = "9.jpeg"
    if not os.path.exists(image_path):
        for ext in ["jpg", "png", "JPG"]:
            if os.path.exists(f"9.{ext}"): image_path = f"9.{ext}"; break
    if os.path.exists(image_path): st.image(image_path, use_container_width=True)
    st.markdown("<h3 style='text-align: center; color: #FF4B4B;'>Mal mal bakma ekrana dosya indi indirilenlere bak</h3>", unsafe_allow_html=True)
    if st.button("Tamam"): st.session_state['popup_ac'] = False; st.rerun()

if st.session_state['popup_ac']: show_popup_modal()

